"""LinkedIn social posting endpoints — settings, AI trends/post/image generation, publish, history."""

import base64
import json
import uuid

import httpx
from fastapi import APIRouter, Request

from routers.config import load_app_config, save_app_config

router = APIRouter(prefix="/api/social", tags=["social"])

LINKEDIN_SETTING_KEYS = ("client_id", "client_secret", "access_token", "user_id", "custom_prompts", "image_model")
IMAGE_MODELS = ("dall-e-3", "gpt-image-1", "gpt-5.5", "dall-e-2")  # tried in order if preferred fails

BUILTIN_TEMPLATES = [
    {
        "id": "builtin-professional-corporate",
        "name": "Professional Corporate",
        "icon": "🏢",
        "prompt": "Clean corporate boardroom or modern office setting, professional lighting, sleek architecture, muted blues and grays, no text overlays, business LinkedIn aesthetic",
        "sample_image": "",
        "builtin": 1,
    },
    {
        "id": "builtin-inspirational-quote",
        "name": "Inspirational Quote",
        "icon": "💬",
        "prompt": "Minimalist motivational poster style, bold clean typography placeholder on a smooth gradient background, warm amber or teal tones, uplifting atmosphere, no specific text",
        "sample_image": "",
        "builtin": 1,
    },
    {
        "id": "builtin-tech-innovation",
        "name": "Tech & Innovation",
        "icon": "🚀",
        "prompt": "Futuristic technology visualization, blue holographic UI elements, glowing circuit patterns, data flow streams, dark background with neon accents, high-tech professional aesthetic",
        "sample_image": "",
        "builtin": 1,
    },
    {
        "id": "builtin-warm-storytelling",
        "name": "Warm & Storytelling",
        "icon": "🌟",
        "prompt": "Authentic candid moment, warm golden natural lighting, genuine human connection or team collaboration, documentary photography style, relatable and approachable, no artificial staging",
        "sample_image": "",
        "builtin": 1,
    },
    {
        "id": "builtin-data-analytics",
        "name": "Data & Analytics",
        "icon": "📊",
        "prompt": "Clean business data visualization, professional charts and graphs on a white or light blue background, modern dashboard aesthetic, crisp lines, no specific numbers or labels",
        "sample_image": "",
        "builtin": 1,
    },
    {
        "id": "builtin-leadership-growth",
        "name": "Leadership & Growth",
        "icon": "📈",
        "prompt": "Upward growth and success visualization, staircase to achievement or mountain summit metaphor, inspiring greens and golds, teamwork or individual triumph, professional motivational imagery",
        "sample_image": "",
        "builtin": 1,
    },
]


def _ensure_tables(cache):
    with cache._conn() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS linkedin_posts (
                id TEXT PRIMARY KEY,
                subject TEXT,
                topic TEXT,
                post_text TEXT,
                audience TEXT,
                tone TEXT,
                image_url TEXT,
                image_prompt TEXT,
                content_type TEXT,
                scheduled_at TEXT,
                published_at TEXT,
                linkedin_post_id TEXT,
                status TEXT DEFAULT 'draft',
                created_at TEXT DEFAULT (datetime('now'))
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS linkedin_templates (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                prompt TEXT NOT NULL,
                sample_image TEXT DEFAULT '',
                builtin INTEGER DEFAULT 0,
                created_at TEXT DEFAULT (datetime('now'))
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS linkedin_autopilot (
                id INTEGER PRIMARY KEY,
                topics TEXT NOT NULL DEFAULT '[]',
                template_id TEXT,
                template_prompt TEXT,
                content_type TEXT DEFAULT 'image+text',
                interval_days INTEGER DEFAULT 7,
                post_time TEXT DEFAULT '09:00',
                enabled INTEGER DEFAULT 1,
                topic_index INTEGER DEFAULT 0,
                last_post_at TEXT,
                next_post_at TEXT,
                require_review INTEGER DEFAULT 0,
                fixed_hashtags TEXT DEFAULT '[]',
                created_at TEXT DEFAULT (datetime('now'))
            )
        """)
        # Migrate existing rows — safe no-op if column already exists
        for col, defval in [("require_review", "0"), ("fixed_hashtags", "'[]'")]:
            try:
                conn.execute(f"ALTER TABLE linkedin_autopilot ADD COLUMN {col} INTEGER DEFAULT {defval}")
            except Exception:
                pass


def _get_linkedin_settings() -> dict:
    cfg = load_app_config()
    ln = cfg.get("linkedin", {}) or {}
    return {
        "client_id": ln.get("client_id", ""),
        "client_secret": ln.get("client_secret", ""),
        "access_token": ln.get("access_token", ""),
        "user_id": ln.get("user_id", ""),
        "custom_prompts": ln.get("custom_prompts", []),
        "image_model": ln.get("image_model", "dall-e-3"),
    }


def _get_openai_key() -> str:
    cfg = load_app_config()
    # Check ai_providers first — it's the actively managed list and always up-to-date
    for p in cfg.get("ai_providers", []):
        if p.get("type") == "openai" and p.get("key"):
            return p["key"]
    return cfg.get("openai_api_key", "")


async def _ai_complete(request: Request, prompt: str, max_tokens: int = 800, system: str = "") -> str:
    advisor = getattr(request.app.state, "advisor", None)
    if not advisor:
        raise RuntimeError("AI advisor not available")
    messages = [{"role": "user", "content": prompt}]
    ant = getattr(advisor.ai, "_anthropic", None)
    kwargs: dict = {"model": "claude-opus-4-8", "max_tokens": max_tokens, "messages": messages}
    if system:
        kwargs["system"] = system
    if ant:
        resp = await ant.messages.create(**kwargs)
    else:
        resp = await advisor.ai.messages.create(**kwargs)
    return resp.content[0].text


def _extract_json(text: str):
    """Pull a JSON array/object out of an AI response that may wrap it in prose or code fences."""
    import re
    text = text.strip()
    # Strip markdown code fences
    if text.startswith("```"):
        inner = re.sub(r"^```[a-z]*\n?", "", text)
        inner = re.sub(r"```$", "", inner).strip()
        if inner:
            text = inner
    # Try direct parse first
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    # Find outermost { } or [ ]
    for open_ch, close_ch in (("{", "}"), ("[", "]")):
        start = text.find(open_ch)
        end = text.rfind(close_ch)
        if start != -1 and end != -1 and end > start:
            chunk = text[start:end + 1]
            try:
                return json.loads(chunk)
            except json.JSONDecodeError:
                # Try replacing single-quote keys/values (AI sometimes uses them)
                try:
                    import ast
                    return ast.literal_eval(chunk)
                except Exception:
                    pass
    return None


@router.get("/linkedin/settings")
async def get_settings(request: Request):
    return _get_linkedin_settings()


@router.post("/linkedin/settings")
async def save_settings(body: dict, request: Request):
    cfg = load_app_config()
    ln = cfg.get("linkedin", {}) or {}
    for key in LINKEDIN_SETTING_KEYS:
        if key in body and body[key] is not None:
            ln[key] = body[key]
    cfg["linkedin"] = ln
    save_app_config(cfg)
    return {"status": "saved"}


@router.post("/linkedin/trends")
async def get_trends(body: dict, request: Request):
    subject = (body.get("subject") or "").strip()
    if not subject:
        return {"trends": [], "error": "subject is required"}
    prompt = (
        f"Generate 5 current trending LinkedIn post topics related to: {subject}. "
        "Return as JSON array: "
        "[{'title': '...', 'description': '...', 'engagement': 'High/Medium', 'hashtags': ['...']}]"
    )
    try:
        content = await _ai_complete(request, prompt, max_tokens=800)
    except Exception as e:
        return {"trends": [], "error": str(e)}
    trends = _extract_json(content)
    if not isinstance(trends, list):
        return {"trends": [], "error": "Could not parse AI response"}
    return {"trends": trends}


@router.post("/linkedin/generate-post")
async def generate_post(body: dict, request: Request):
    topic = (body.get("topic") or "").strip()
    audience = (body.get("audience") or "General").strip()
    tone = (body.get("tone") or "Professional").strip()
    subject = (body.get("subject") or "").strip()
    if not topic:
        return {"post": "", "hashtags": [], "char_count": 0, "error": "topic is required"}

    # --- Call 1: get ONLY the post text (no JSON, no markers, pure text) ---
    post_system = (
        "You are a LinkedIn post writer. "
        "Output ONLY the post text — nothing else. "
        "No JSON, no code blocks, no labels, no introductions, no explanations."
    )
    post_prompt = (
        f"Write a professional LinkedIn post about: {topic}.\n"
        f"Subject area: {subject or topic}. Target audience: {audience}. Tone: {tone}.\n"
        "Engaging, well-structured with real line breaks. Emojis sparingly. "
        "Do NOT include hashtags in the body. "
        "Output ONLY the post text — start writing the post immediately."
    )
    try:
        post_raw = await _ai_complete(request, post_prompt, max_tokens=1200, system=post_system)
    except Exception as e:
        return {"post": "", "hashtags": [], "char_count": 0, "error": str(e)}

    post = post_raw.strip()
    # Strip any accidental JSON wrapper (old model behaviour)
    if post.startswith("{") and ('"post"' in post or "'post'" in post):
        parsed = _extract_json(post)
        if isinstance(parsed, dict) and "post" in parsed:
            post = str(parsed["post"]).strip()

    # --- Call 2: get hashtags as a simple comma-separated list ---
    hashtag_system = "Return ONLY a comma-separated list of hashtag words without the # symbol. No JSON. No other text."
    hashtag_prompt = f"Give 6 LinkedIn hashtags for a post about: {topic}. Return ONLY comma-separated words."
    hashtags: list[str] = []
    try:
        ht_raw = await _ai_complete(request, hashtag_prompt, max_tokens=80, system=hashtag_system)
        hashtags = [h.strip().lstrip("#") for h in ht_raw.split(",") if h.strip() and len(h.strip()) < 35]
    except Exception:
        pass

    return {"post": post, "hashtags": hashtags, "char_count": len(post)}


@router.post("/linkedin/generate-images")
async def generate_images(body: dict, request: Request):
    topic = (body.get("topic") or "").strip()
    post_text = (body.get("post_text") or "").strip()
    custom_prompt = body.get("custom_prompt")

    openai_key = _get_openai_key()
    if not openai_key:
        return {"images": [], "error": "OpenAI API key not configured — add it in Settings → AI Providers (type: openai)"}

    if not openai_key.startswith("sk-"):
        return {"images": [], "error": f"OpenAI key looks invalid (should start with 'sk-'). Current key starts with: {openai_key[:8]}…"}

    base = custom_prompt.strip() if custom_prompt else ""
    post_summary = post_text[:600] if post_text else topic

    if base:
        # Template/custom prompt provided: combine directly with post details — no AI rewrite
        topic_line = f"Topic: {topic}." if topic else ""
        summary_line = f" Context: {post_summary[:300]}" if post_summary else ""
        dalle_prompt = f"{base}. {topic_line}{summary_line}".strip(". ").strip()
    else:
        # No template: use AI to write a meaningful DALL-E prompt from scratch
        img_system = "Return ONLY the image prompt text — no JSON, no labels, no extra text."
        img_prompt = (
            "Write ONE detailed DALL-E image prompt for a LinkedIn post.\n\n"
            f"POST TOPIC: {topic}\n"
            f"POST TEXT (use this for the visual theme):\n{post_summary}\n\n"
            "The prompt must:\n"
            "- Visually represent the core message of the post\n"
            "- Be professional and LinkedIn-appropriate\n"
            "- Include composition, lighting, color tone, and mood\n"
            "Return ONLY the prompt text — nothing else."
        )
        try:
            dalle_prompt = await _ai_complete(request, img_prompt, max_tokens=300, system=img_system)
            dalle_prompt = dalle_prompt.strip().strip('"').strip("'")
        except Exception as e:
            return {"images": [], "error": f"AI prompt generation failed: {e}"}

        if not dalle_prompt:
            return {"images": [], "error": "Could not generate image prompt"}

    preferred = _get_linkedin_settings().get("image_model", "dall-e-3") or "dall-e-3"
    fallback_list = list(IMAGE_MODELS)
    if preferred in fallback_list:
        fallback_list.remove(preferred)
    IMAGE_MODEL_FALLBACKS = [preferred] + fallback_list

    headers = {"Authorization": f"Bearer {openai_key}", "Content-Type": "application/json"}

    async def _try_image(http_client, p: str, model: str) -> tuple[str, str]:
        payload: dict = {"model": model, "prompt": p[:4000], "n": 1, "size": "1024x1024"}
        try:
            r = await http_client.post(
                "https://api.openai.com/v1/images/generations", headers=headers, json=payload
            )
        except Exception as e:
            return "", f"Request failed: {e}"
        if r.status_code == 200:
            item = r.json().get("data", [{}])[0]
            url = item.get("url", "")
            b64 = item.get("b64_json", "")
            if url:
                return url, ""
            if b64:
                return f"data:image/png;base64,{b64}", ""
            return "", f"Model {model} returned no image data"
        eb = r.json() if r.headers.get("content-type", "").startswith("application/json") else {}
        msg = eb.get("error", {}).get("message", r.text[:300])
        return "", f"OpenAI {r.status_code}: {msg}"

    last_error = ""
    _skip_phrases = ("does not exist", "not supported", "not available", "no access", "model_not_found")
    async with httpx.AsyncClient(timeout=120.0) as http:
        for m in IMAGE_MODEL_FALLBACKS:
            url, err = await _try_image(http, dalle_prompt, m)
            if url:
                return {"images": [{"url": url, "prompt": dalle_prompt}]}
            last_error = err
            if not any(p in err.lower() for p in _skip_phrases):
                break  # real error — no point trying other models

    return {"images": [], "error": last_error or "Image generation failed"}


@router.post("/linkedin/save-draft")
async def save_draft(body: dict, request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    post_id = str(uuid.uuid4())
    with cache._conn() as conn:
        conn.execute(
            """
            INSERT INTO linkedin_posts
                (id, subject, topic, post_text, audience, tone, image_url, image_prompt,
                 content_type, scheduled_at, status)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'draft')
            """,
            (
                post_id,
                body.get("subject"),
                body.get("topic"),
                body.get("post_text"),
                body.get("audience"),
                body.get("tone"),
                body.get("image_url"),
                body.get("image_prompt"),
                body.get("content_type"),
                body.get("scheduled_at"),
            ),
        )
    return {"id": post_id}


async def _resolve_linkedin_author(access_token: str, stored_user_id: str) -> tuple[str, str]:
    """Returns (urn:li:person:ID, error). Always tries /userinfo first for the correct sub."""
    # Always try userinfo — it gives the exact sub LinkedIn accepts as author
    try:
        async with httpx.AsyncClient(timeout=10.0) as http:
            r = await http.get(
                "https://api.linkedin.com/v2/userinfo",
                headers={"Authorization": f"Bearer {access_token}"},
            )
        if r.status_code == 200:
            sub = r.json().get("sub", "")
            if sub:
                # UGC Posts API requires urn:li:person: (not urn:li:member:)
                return f"urn:li:person:{sub}", ""
    except Exception:
        pass
    # Fall back to stored value
    if stored_user_id:
        if stored_user_id.startswith("urn:li:person:"):
            return stored_user_id, ""
        if stored_user_id.startswith("urn:"):
            return stored_user_id, ""
        return f"urn:li:person:{stored_user_id}", ""
    return "", "LinkedIn User ID not set — go to Settings → LinkedIn"


async def _upload_image_to_linkedin(access_token: str, author: str, image_url: str) -> tuple[str, str]:
    """Upload an image to LinkedIn and return (image_urn, error). image_url can be http(s) or data URI."""
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json",
        "LinkedIn-Version": "202410",
    }
    async with httpx.AsyncClient(timeout=60.0) as http:
        # 1. Initialize upload
        init = await http.post(
            "https://api.linkedin.com/v2/images?action=initializeUpload",
            headers=headers,
            json={"initializeUploadRequest": {"owner": author}},
        )
        if init.status_code >= 400:
            return "", f"Image upload init failed {init.status_code}: {init.text[:200]}"
        val = init.json().get("value", {})
        upload_url = val.get("uploadUrl", "")
        image_urn = val.get("image", "")
        if not upload_url or not image_urn:
            return "", "LinkedIn did not return upload URL"

        # 2. Get image bytes
        if image_url.startswith("data:"):
            _, b64data = image_url.split(",", 1)
            image_bytes = base64.b64decode(b64data)
        else:
            dl = await http.get(image_url, follow_redirects=True)
            if dl.status_code >= 400:
                return "", f"Could not download image: {dl.status_code}"
            image_bytes = dl.content

        # 3. Upload binary — LinkedIn pre-signed URL requires application/octet-stream
        up = await http.put(upload_url, content=image_bytes, headers={"Content-Type": "application/octet-stream"})
        if up.status_code >= 400:
            return "", f"Image binary upload failed {up.status_code}: {up.text[:100]}"

    return image_urn, ""


async def _publish_to_linkedin(post_text: str, settings: dict, image_url: str = "", content_type: str = "image+text") -> dict:
    access_token = settings.get("access_token", "")
    user_id = settings.get("user_id", "")
    if not access_token:
        return {"error": "LinkedIn access token not configured — go to Settings → LinkedIn"}

    author, err = await _resolve_linkedin_author(access_token, user_id)
    if not author:
        return {"error": err or "Could not resolve LinkedIn author ID"}

    api_headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json",
        "X-Restli-Protocol-Version": "2.0.0",
        "LinkedIn-Version": "202410",
    }

    # For image-only posts, commentary is empty
    commentary = "" if content_type == "image" else post_text

    # Upload image if provided
    image_urn = ""
    img_upload_error = ""
    if image_url:
        image_urn, img_upload_error = await _upload_image_to_linkedin(access_token, author, image_url)
        # If upload failed, always surface the error — never silently post text-only
        # when the caller explicitly requested an image.
        if img_upload_error:
            return {"error": f"Image upload failed: {img_upload_error}"}

    # Build new /v2/posts payload
    new_payload: dict = {
        "author": author,
        "commentary": commentary,
        "visibility": "PUBLIC",
        "distribution": {
            "feedDistribution": "MAIN_FEED",
            "targetEntities": [],
            "thirdPartyDistributionChannels": [],
        },
        "lifecycleState": "PUBLISHED",
        "isReshareDisabledByAuthor": False,
    }
    if image_urn:
        new_payload["content"] = {"media": {"id": image_urn}}

    async with httpx.AsyncClient(timeout=30.0) as http:
        resp = await http.post(
            "https://api.linkedin.com/v2/posts", headers=api_headers, json=new_payload
        )

    # If new API failed, try legacy ugcPosts — include image_urn if available
    if resp.status_code in (400, 403, 422):
        if image_urn:
            share_content: dict = {
                "shareCommentary": {"text": post_text},
                "shareMediaCategory": "IMAGE",
                "media": [{"status": "READY", "media": image_urn}],
            }
        else:
            share_content = {
                "shareCommentary": {"text": post_text},
                "shareMediaCategory": "NONE",
            }
        legacy_payload = {
            "author": author,
            "lifecycleState": "PUBLISHED",
            "specificContent": {"com.linkedin.ugc.ShareContent": share_content},
            "visibility": {"com.linkedin.ugc.MemberNetworkVisibility": "PUBLIC"},
        }
        async with httpx.AsyncClient(timeout=30.0) as http:
            resp2 = await http.post(
                "https://api.linkedin.com/v2/ugcPosts",
                headers={k: v for k, v in api_headers.items() if k != "LinkedIn-Version"},
                json=legacy_payload,
            )
        if resp2.status_code < 400:
            resp = resp2
        elif resp.status_code >= 400:
            resp = resp2 if resp2.status_code != 403 else resp

    if resp.status_code >= 400:
        raw = resp.text[:400]
        if resp.status_code == 403:
            return {"error": (
                f"LinkedIn 403: token missing 'w_member_social' scope. "
                f"Author: {author}. "
                f"Fix: in your LinkedIn app → Products enable 'Share on LinkedIn', then regenerate your access token. "
                f"Raw: {raw}"
            )}
        return {"error": f"LinkedIn API error {resp.status_code}: {raw}"}

    post_id = resp.headers.get("x-restli-id", "") or (resp.json() or {}).get("id", "")
    return {"linkedin_post_id": post_id, "image_uploaded": bool(image_urn)}


@router.post("/linkedin/publish")
async def publish(body: dict, request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    post_id = body.get("id") or str(uuid.uuid4())
    post_text = body.get("post_text") or ""
    scheduled_at = body.get("scheduled_at")
    content_type = body.get("content_type", "image+text")
    image_url = body.get("image_url") or ""
    topic = (body.get("topic") or "").strip()
    subject = (body.get("subject") or "").strip()

    # Ensure a history record exists (wizard publishes without a prior draft save)
    with cache._conn() as conn:
        exists = conn.execute(
            "SELECT 1 FROM linkedin_posts WHERE id=?", (post_id,)
        ).fetchone()
        if not exists:
            conn.execute(
                """INSERT INTO linkedin_posts
                   (id, topic, subject, post_text, image_url, content_type, status)
                   VALUES (?, ?, ?, ?, ?, ?, 'pending')""",
                (post_id, topic, subject, post_text, image_url, content_type),
            )

    if scheduled_at:
        with cache._conn() as conn:
            conn.execute(
                "UPDATE linkedin_posts SET status='scheduled', scheduled_at=? WHERE id=?",
                (scheduled_at, post_id),
            )
        return {"status": "scheduled", "scheduled_for": scheduled_at}

    settings = _get_linkedin_settings()
    result = await _publish_to_linkedin(post_text, settings, image_url, content_type)
    if "error" in result:
        with cache._conn() as conn:
            conn.execute(
                "UPDATE linkedin_posts SET status='failed' WHERE id=?", (post_id,)
            )
        return {"error": result["error"]}

    linkedin_post_id = result.get("linkedin_post_id", "")
    with cache._conn() as conn:
        conn.execute(
            """UPDATE linkedin_posts
               SET status='published', published_at=datetime('now'), linkedin_post_id=?,
                   image_url=COALESCE(NULLIF(?, ''), image_url)
               WHERE id=?""",
            (linkedin_post_id, image_url, post_id),
        )
    return {"status": "published", "linkedin_post_id": linkedin_post_id, "image_uploaded": result.get("image_uploaded", False)}


@router.get("/linkedin/history")
async def get_history(request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        rows = conn.execute(
            "SELECT * FROM linkedin_posts ORDER BY created_at DESC"
        ).fetchall()
    return {"posts": [dict(r) for r in rows]}


@router.delete("/linkedin/history/{post_id}")
async def delete_history(post_id: str, request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        conn.execute("DELETE FROM linkedin_posts WHERE id=?", (post_id,))
    return {"status": "deleted"}


@router.post("/linkedin/history/{post_id}/reschedule")
async def reschedule_post(post_id: str, body: dict, request: Request):
    """Set a new scheduled_at time and reset status to 'scheduled'."""
    cache = request.app.state.cache
    _ensure_tables(cache)
    scheduled_at = (body.get("scheduled_at") or "").strip()
    if not scheduled_at:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail="scheduled_at is required")
    with cache._conn() as conn:
        conn.execute(
            "UPDATE linkedin_posts SET status='scheduled', scheduled_at=? WHERE id=?",
            (scheduled_at, post_id),
        )
    return {"status": "scheduled", "scheduled_for": scheduled_at}


@router.post("/linkedin/ask")
async def ask_about_post(body: dict, request: Request):
    """Ask an AI question about a past post topic."""
    topic = (body.get("topic") or "").strip()
    post_text = (body.get("post_text") or "").strip()
    question = (body.get("question") or "").strip()
    if not question:
        return {"error": "question is required"}
    system = (
        "You are a LinkedIn content strategist. Answer concisely and helpfully. "
        "No JSON, no markdown headers — plain text only."
    )
    context = ""
    if topic:
        context += f"Post topic: {topic}\n"
    if post_text:
        context += f"Post text (excerpt): {post_text[:500]}\n"
    prompt = f"{context}\nQuestion: {question}"
    answer = await _ai_complete(request, prompt, max_tokens=400, system=system)
    return {"answer": answer.strip()}


# ── Prompt Template Library ────────────────────────────────────────────────────

@router.get("/linkedin/templates")
async def get_templates(request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        rows = conn.execute(
            "SELECT id, name, prompt, sample_image, builtin, created_at FROM linkedin_templates WHERE builtin=0 ORDER BY created_at ASC"
        ).fetchall()
    user_templates = [dict(r) for r in rows]
    return {"templates": BUILTIN_TEMPLATES + user_templates}


@router.post("/linkedin/templates")
async def save_template(body: dict, request: Request):
    name = (body.get("name") or "").strip()
    prompt = (body.get("prompt") or "").strip()
    if not name or not prompt:
        return {"error": "name and prompt are required"}
    cache = request.app.state.cache
    _ensure_tables(cache)
    tmpl_id = str(uuid.uuid4())
    sample_image = (body.get("sample_image") or "")
    with cache._conn() as conn:
        conn.execute(
            "INSERT INTO linkedin_templates (id, name, prompt, sample_image, builtin) VALUES (?, ?, ?, ?, 0)",
            (tmpl_id, name, prompt, sample_image),
        )
    return {"id": tmpl_id, "name": name, "prompt": prompt, "sample_image": sample_image, "builtin": 0}


@router.put("/linkedin/templates/{template_id}")
async def update_template(template_id: str, body: dict, request: Request):
    name = (body.get("name") or "").strip()
    prompt = (body.get("prompt") or "").strip()
    if not name or not prompt:
        return {"error": "name and prompt are required"}
    cache = request.app.state.cache
    _ensure_tables(cache)
    sample_image = body.get("sample_image")
    with cache._conn() as conn:
        if sample_image is not None:
            conn.execute(
                "UPDATE linkedin_templates SET name=?, prompt=?, sample_image=? WHERE id=? AND builtin=0",
                (name, prompt, sample_image, template_id),
            )
        else:
            conn.execute(
                "UPDATE linkedin_templates SET name=?, prompt=? WHERE id=? AND builtin=0",
                (name, prompt, template_id),
            )
    return {"status": "updated"}


@router.delete("/linkedin/templates/{template_id}")
async def delete_template(template_id: str, request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        conn.execute("DELETE FROM linkedin_templates WHERE id=? AND builtin=0", (template_id,))
    return {"status": "deleted"}


@router.post("/linkedin/templates/improve-prompt")
async def improve_template_prompt(body: dict, request: Request):
    """Use AI to rewrite/improve a DALL-E image prompt for a LinkedIn template."""
    current_prompt = (body.get("prompt") or "").strip()
    instruction = (body.get("instruction") or "").strip()
    if not current_prompt:
        return {"error": "prompt is required"}

    system = (
        "You are an expert DALL-E prompt engineer specializing in LinkedIn professional imagery. "
        "You rewrite image-style prompts to be more vivid, specific, and effective for AI image generation. "
        "Keep prompts concise (2-4 sentences). Do NOT include text, words, or letters in the image description. "
        "Return ONLY the improved prompt — no explanation, no preamble."
    )
    user_msg = f"Current prompt:\n{current_prompt}"
    if instruction:
        user_msg += f"\n\nImprovement instruction: {instruction}"
    else:
        user_msg += "\n\nImprove this prompt to be more specific, evocative, and effective for DALL-E."

    advisor = getattr(request.app.state, "advisor", None)
    if not advisor:
        return {"error": "AI advisor not available"}
    ant = getattr(advisor.ai, "_anthropic", None)
    kwargs = dict(model="claude-haiku-4-5-20251001", max_tokens=400, system=system,
                  messages=[{"role": "user", "content": user_msg}])
    try:
        if ant:
            resp = await ant.messages.create(**kwargs)
        else:
            resp = await advisor.ai.messages.create(**kwargs)
        improved = resp.content[0].text.strip()
        return {"improved_prompt": improved}
    except Exception as e:
        return {"error": str(e)}


# ── Connectivity Verification ──────────────────────────────────────────────────

@router.post("/linkedin/verify")
async def verify_linkedin(request: Request):
    """Verify LinkedIn token, OpenAI key, and AI provider connectivity."""
    results = {}

    # LinkedIn
    settings = _get_linkedin_settings()
    access_token = settings.get("access_token", "")
    if not access_token:
        results["linkedin"] = {"ok": False, "message": "No access token configured"}
    else:
        try:
            async with httpx.AsyncClient(timeout=10.0) as http:
                r = await http.get(
                    "https://api.linkedin.com/v2/userinfo",
                    headers={"Authorization": f"Bearer {access_token}"},
                )
            if r.status_code == 200:
                data = r.json()
                name = data.get("name") or data.get("given_name", "")
                sub = data.get("sub", "")
                author_urn = f"urn:li:person:{sub}" if sub else "(sub missing)"
                results["linkedin"] = {
                    "ok": True,
                    "message": f"Connected as {name} · author URN: {author_urn}" if name else f"Connected · {author_urn}",
                }
            else:
                results["linkedin"] = {"ok": False, "message": f"HTTP {r.status_code}: token may be expired"}
        except Exception as e:
            results["linkedin"] = {"ok": False, "message": str(e)}

    # OpenAI
    openai_key = _get_openai_key()
    if not openai_key:
        results["openai"] = {"ok": False, "message": "No OpenAI API key configured"}
    else:
        try:
            async with httpx.AsyncClient(timeout=10.0) as http:
                r = await http.get(
                    "https://api.openai.com/v1/models",
                    headers={"Authorization": f"Bearer {openai_key}"},
                )
            if r.status_code == 200:
                results["openai"] = {"ok": True, "message": "OpenAI API key valid"}
            else:
                results["openai"] = {"ok": False, "message": f"HTTP {r.status_code}: invalid key"}
        except Exception as e:
            results["openai"] = {"ok": False, "message": str(e)}

    # AI provider (Claude/Anthropic)
    try:
        advisor = getattr(request.app.state, "advisor", None)
        if advisor:
            await _ai_complete(request, "Reply with the single word: ok", max_tokens=10)
            results["ai_provider"] = {"ok": True, "message": "AI provider responding"}
        else:
            results["ai_provider"] = {"ok": False, "message": "AI advisor not initialised"}
    except Exception as e:
        results["ai_provider"] = {"ok": False, "message": str(e)}

    return results


# ── LinkedIn Autopilot ─────────────────────────────────────────────────────────

@router.get("/linkedin/autopilot")
async def get_autopilot(request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        row = conn.execute("SELECT * FROM linkedin_autopilot ORDER BY id LIMIT 1").fetchone()
    if not row:
        return {"config": None}
    d = dict(row)
    d["topics"] = json.loads(d.get("topics") or "[]")
    d["fixed_hashtags"] = json.loads(d.get("fixed_hashtags") or "[]")
    return {"config": d}


@router.post("/linkedin/autopilot")
async def save_autopilot(body: dict, request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    topics = body.get("topics") or []
    template_id = body.get("template_id") or None
    content_type = body.get("content_type", "image+text")
    interval_days = int(body.get("interval_days") or 7)
    post_time = body.get("post_time") or "09:00"
    enabled = 1 if body.get("enabled", True) else 0
    next_post_at = body.get("next_post_at") or None
    topic_index = int(body.get("topic_index") or 0)
    require_review = 1 if body.get("require_review") else 0
    fixed_hashtags = json.dumps(body.get("fixed_hashtags") or [])

    # Fetch template prompt if a template is selected
    template_prompt = ""
    if template_id:
        with cache._conn() as conn:
            t = conn.execute("SELECT prompt FROM linkedin_templates WHERE id=?", (template_id,)).fetchone()
            if t:
                template_prompt = t["prompt"]

    topics_json = json.dumps(topics)
    with cache._conn() as conn:
        existing = conn.execute("SELECT id FROM linkedin_autopilot LIMIT 1").fetchone()
        if existing:
            conn.execute(
                """UPDATE linkedin_autopilot SET topics=?, template_id=?, template_prompt=?,
                   content_type=?, interval_days=?, post_time=?, enabled=?, next_post_at=?,
                   topic_index=?, require_review=?, fixed_hashtags=? WHERE id=?""",
                (topics_json, template_id, template_prompt, content_type, interval_days,
                 post_time, enabled, next_post_at, topic_index, require_review, fixed_hashtags,
                 existing["id"]),
            )
        else:
            conn.execute(
                """INSERT INTO linkedin_autopilot
                   (topics, template_id, template_prompt, content_type, interval_days,
                    post_time, enabled, next_post_at, topic_index, require_review, fixed_hashtags)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (topics_json, template_id, template_prompt, content_type, interval_days,
                 post_time, enabled, next_post_at, topic_index, require_review, fixed_hashtags),
            )
    return {"status": "saved"}


@router.delete("/linkedin/autopilot")
async def delete_autopilot(request: Request):
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        conn.execute("DELETE FROM linkedin_autopilot")
    return {"status": "deleted"}


@router.get("/linkedin/autopilot/review-queue")
async def get_review_queue(request: Request):
    """Get posts awaiting approval before publishing."""
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        rows = conn.execute(
            "SELECT * FROM linkedin_posts WHERE status='pending_review' ORDER BY created_at DESC"
        ).fetchall()
    return {"posts": [dict(r) for r in rows]}


@router.post("/linkedin/autopilot/review/{post_id}/approve")
async def approve_review_post(post_id: str, request: Request):
    """Publish a pending-review post immediately."""
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        row = conn.execute("SELECT * FROM linkedin_posts WHERE id=?", (post_id,)).fetchone()
    if not row:
        from fastapi import HTTPException
        raise HTTPException(404, "Post not found")
    post = dict(row)
    settings = _get_linkedin_settings()
    result = await _publish_to_linkedin(
        post["post_text"], settings,
        post.get("image_url") or "", post.get("content_type") or "image+text"
    )
    if "error" in result:
        return {"error": result["error"]}
    linkedin_post_id = result.get("linkedin_post_id", "")
    with cache._conn() as conn:
        conn.execute(
            "UPDATE linkedin_posts SET status='published', published_at=datetime('now'), linkedin_post_id=? WHERE id=?",
            (linkedin_post_id, post_id),
        )
    return {"status": "published", "linkedin_post_id": linkedin_post_id}


@router.post("/linkedin/autopilot/review/{post_id}/edit-approve")
async def edit_approve_review_post(post_id: str, body: dict, request: Request):
    """Update post text then publish."""
    cache = request.app.state.cache
    _ensure_tables(cache)
    new_text = (body.get("post_text") or "").strip()
    if not new_text:
        from fastapi import HTTPException
        raise HTTPException(400, "post_text is required")
    with cache._conn() as conn:
        conn.execute("UPDATE linkedin_posts SET post_text=? WHERE id=?", (new_text, post_id))
        row = conn.execute("SELECT * FROM linkedin_posts WHERE id=?", (post_id,)).fetchone()
    if not row:
        from fastapi import HTTPException
        raise HTTPException(404, "Post not found")
    post = dict(row)
    settings = _get_linkedin_settings()
    result = await _publish_to_linkedin(
        new_text, settings,
        post.get("image_url") or "", post.get("content_type") or "image+text"
    )
    if "error" in result:
        return {"error": result["error"]}
    linkedin_post_id = result.get("linkedin_post_id", "")
    with cache._conn() as conn:
        conn.execute(
            "UPDATE linkedin_posts SET status='published', published_at=datetime('now'), linkedin_post_id=? WHERE id=?",
            (linkedin_post_id, post_id),
        )
    return {"status": "published", "linkedin_post_id": linkedin_post_id}


@router.post("/linkedin/autopilot/review/{post_id}/reject")
async def reject_review_post(post_id: str, request: Request):
    """Reject and discard a pending-review post."""
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        conn.execute("UPDATE linkedin_posts SET status='rejected' WHERE id=?", (post_id,))
    return {"status": "rejected"}


@router.get("/linkedin/history/{post_id}/stats")
async def get_post_stats(post_id: str, request: Request):
    """Return LinkedIn post URL for viewing analytics in-browser (API analytics require LinkedIn partnership)."""
    cache = request.app.state.cache
    _ensure_tables(cache)
    with cache._conn() as conn:
        row = conn.execute("SELECT linkedin_post_id FROM linkedin_posts WHERE id=?", (post_id,)).fetchone()
    if not row or not row["linkedin_post_id"]:
        return {"linkedin_url": None, "api_available": False}
    linkedin_urn = row["linkedin_post_id"]
    # Build the direct LinkedIn post URL
    linkedin_url = f"https://www.linkedin.com/feed/update/{linkedin_urn}"
    return {"linkedin_url": linkedin_url, "api_available": False}


@router.post("/linkedin/autopilot/extract-topics")
async def extract_topics_from_file(request: Request):
    """Upload a resume/PDF/DOCX and extract LinkedIn post topics from it using AI."""
    import tempfile, os
    from fastapi import UploadFile
    form = await request.form()
    file: UploadFile = form.get("file")  # type: ignore
    if not file:
        return {"topics": [], "error": "No file uploaded"}

    suffix = os.path.splitext(file.filename or "file.txt")[1].lower() or ".txt"
    if suffix not in (".pdf", ".docx", ".txt", ".md", ".rtf"):
        return {"topics": [], "error": f"Unsupported file type: {suffix}. Use PDF, DOCX, or TXT."}

    content = await file.read()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    import subprocess, glob as _glob, shutil

    text = ""
    extract_error = ""
    try:
        if suffix == ".pdf":
            # Tier 1: pdfminer (text-based PDFs)
            from pdfminer.high_level import extract_text as _pdf_extract
            text = (_pdf_extract(tmp_path) or "").strip()

            # Tier 2: pdftotext from poppler (handles more PDF variants)
            if len(text) < 30:
                pt = shutil.which("pdftotext") or "/opt/homebrew/bin/pdftotext"
                if os.path.exists(pt):
                    r = subprocess.run([pt, tmp_path, "-"], capture_output=True, timeout=30)
                    text = (r.stdout.decode("utf-8", errors="replace") or "").strip()

            # Tier 3: OCR via pdftoppm + tesseract (scanned/image PDFs)
            if len(text) < 30:
                pdftoppm = shutil.which("pdftoppm") or "/opt/homebrew/bin/pdftoppm"
                tesseract = shutil.which("tesseract") or "/opt/homebrew/bin/tesseract"
                if os.path.exists(pdftoppm) and os.path.exists(tesseract):
                    import tempfile as _tf
                    with _tf.TemporaryDirectory() as ocr_dir:
                        prefix = os.path.join(ocr_dir, "page")
                        subprocess.run(
                            [pdftoppm, "-r", "150", "-png", "-l", "4", tmp_path, prefix],
                            capture_output=True, timeout=60
                        )
                        parts = []
                        for img in sorted(_glob.glob(prefix + "-*.png"))[:4]:
                            r2 = subprocess.run(
                                [tesseract, img, "stdout", "-l", "eng"],
                                capture_output=True, timeout=30
                            )
                            parts.append(r2.stdout.decode("utf-8", errors="replace"))
                        text = "\n".join(parts).strip()

        elif suffix == ".docx":
            import docx as _docx
            doc = _docx.Document(tmp_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            text = "\n".join(parts)
        else:
            text = open(tmp_path, encoding="utf-8", errors="replace").read()
    except Exception as e:
        extract_error = str(e)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    if extract_error:
        return {"topics": [], "error": f"Could not read file: {extract_error}"}
    if not text or len(text.strip()) < 30:
        return {"topics": [], "error": "File appears empty or unreadable. Ensure the PDF is not password-protected."}

    prompt = (
        "You are a LinkedIn content strategist. Analyze the following document (resume, bio, or profile) "
        "and generate 10 highly relevant LinkedIn post topics this person should write about based on their "
        "expertise, experience, skills, and industry.\n\n"
        "Rules:\n"
        "- Each topic must be specific and engaging (not generic like 'Leadership')\n"
        "- Topics should leverage their actual skills/experience from the document\n"
        "- Mix thought leadership, lessons learned, how-tos, and personal stories\n"
        "- Return ONLY a JSON array of strings, e.g.: [\"Topic 1\", \"Topic 2\", ...]\n\n"
        f"DOCUMENT:\n{text[:4000]}"
    )
    try:
        content_text = await _ai_complete(request, prompt, max_tokens=600)
        topics = _extract_json(content_text)
        if not isinstance(topics, list):
            return {"topics": [], "error": "AI could not parse topics from document"}
        topics = [str(t).strip() for t in topics if str(t).strip()][:15]
        return {"topics": topics}
    except Exception as e:
        return {"topics": [], "error": str(e)}
